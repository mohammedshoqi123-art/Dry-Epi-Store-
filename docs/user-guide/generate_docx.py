from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return h

def add_info_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"💡 {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

def add_warn_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"⚠️ {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

def add_success_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"✅ {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

def add_step(num, text):
    p = doc.add_paragraph()
    run = p.add_run(f"الخطوة {num}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    p.add_run(text)

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx+1].cells[c_idx].text = val
    doc.add_paragraph()

# ═══ COVER PAGE ═══
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("💉")
run.font.size = Pt(48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("دليل استخدام منصة مشرف EPI")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EPI Supervisor Platform — User Guide")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("الإصدار 1.0.0 | أبريل 2026")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("نظام إشراف ميداني متكامل لحملات التطعيم")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══ TOC ═══
add_heading_styled("📑 فهرس المحتويات", 1)
toc_items = [
    "1. نظرة عامة على المنصة",
    "2. تسجيل الدخول والحساب",
    "3. لوحة التحكم الرئيسية",
    "4. ملء النماذج الميدانية",
    "5. عرض الإرساليات",
    "6. لوحة التحليلات",
    "7. المساعد الذكي AI",
    "8. الخريطة التفاعلية",
    "9. إدارة المستخدمين (للمدير)",
    "10. إدارة النماذج (للمدير)",
    "11. العمل بدون إنترنت",
    "12. نظام الصلاحيات",
    "13. الإشعارات",
    "14. استكشاف الأخطاء وحلها",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ═══ SECTION 1 ═══
add_heading_styled("1. نظرة عامة على المنصة", 1)
doc.add_paragraph("منصة مشرف EPI هي نظام متكامل لإدارة والإشراف على حملات التطعيم الميدانية. تم تصميمها خصيصاً للعمل في بيئات قد تكون فيها الاتصال بالإنترنت محدوداً أو غير متاح.")

add_heading_styled("المميزات الرئيسية", 2)
add_table(
    ["الميزة", "الوصف"],
    [
        ["📝 نماذج ديناميكية", "محرك JSON Schema قابل للتخصيص"],
        ["📡 عمل بدون إنترنت", "املأ النماذج واحفظها محلياً — تُزامن تلقائياً"],
        ["📊 تحليلات فورية", "KPIs ورسوم بيانية وتقارير PDF/CSV"],
        ["🤖 ذكاء اصطناعي", "مساعد MiMo للتحليل باللغة العربية"],
        ["🗺️ خرائط تفاعلية", "تتبع المواقع مع clustering و heatmap"],
        ["🔐 أمان متقدم", "5 مستويات صلاحيات + تشفير AES-256"],
    ]
)

add_heading_styled("الفئات المستهدفة", 2)
add_step(1, "مدخلو البيانات — ملء النماذج الميدانية وإرسال التقارير")
add_step(2, "مشرفو المديريات — مراجعة الإرساليات ومراقبة الأداء")
add_step(3, "مشرفو المحافظات — الإشراف على جميع مديريات المحافظة")
add_step(4, "المشرف المركزي — مراقبة الأداء على مستوى الجمهورية")
add_step(5, "مدير النظام — إدارة المستخدمين والنماذج والإعدادات")

doc.add_page_break()

# ═══ SECTION 2 ═══
add_heading_styled("2. تسجيل الدخول والحساب", 1)
doc.add_paragraph("يتم تسجيل الدخول باستخدام البريد الإلكتروني وكلمة المرور التي يوفرها مدير النظام.")

add_step(1, "افتح التطبيق — ستظهر شاشة الترحيب ثم تنتقل لشاشة تسجيل الدخول")
add_step(2, "أدخل البريد الإلكتروني الذي سجّل لك من قبل مدير النظام")
add_step(3, "أدخل كلمة المرور واضغط على 'تسجيل الدخول'")
add_step(4, "ستنتقل تلقائياً للوحة التحكم حسب صلاحياتك")

add_warn_box("إذا نسيت كلمة المرور، تواصل مع مدير النظام لإعادة تعيينها.")

doc.add_page_break()

# ═══ SECTION 3 ═══
add_heading_styled("3. لوحة التحكم الرئيسية", 1)
doc.add_paragraph("بعد تسجيل الدخول، تظهر لوحة التحكم التي تعرض ملخص أداء حملات التطعيم.")
doc.add_paragraph("• البطاقات الإحصائية العلوية: تعرض أعداد الإرساليات اليوم ونسبة الإنجاز")
doc.add_paragraph("• الرسوم البيانية: توضح توزيع التطعيمات حسب المحافظة")
doc.add_paragraph("• قائمة التنقل السفلية: تتيح الانتقال السريع بين أقسام التطبيق")

doc.add_page_break()

# ═══ SECTION 4 ═══
add_heading_styled("4. ملء النماذج الميدانية", 1)
doc.add_paragraph("هذه الوظيفة الأساسية للمنصة — ملء نماذج التطعيم الميدانية وإرسالها.")

add_step(1, "افتح قسم 'النماذج' من القائمة السفلية أو الجانبية")
add_step(2, "اختر النموذج المطلوب من القائمة")
add_step(3, "املأ الحقول — النموذج مقسم لأقسام (بيانات الطفل، نوع التطعيم، الموقع)")
add_step(4, "أرفق الصور إن وُجدت (التقط من الكاميرا أو اختر من المعرض)")
add_step(5, "راجع البيانات وتأكد من صحتها")
add_step(6, "اضغط إرسال — يُحفظ محلياً أولاً ثم يُرسل للسيرفر")

add_success_box("البيانات تُحفظ فوراً على الجهاز حتى بدون إنترنت. لن تفقد أي بيانات!")

add_info_box("يمكنك الضغط على 'حفظ كمسودة' إذا لم تنتهِ من الملء.")

doc.add_page_break()

# ═══ SECTION 5 ═══
add_heading_styled("5. عرض الإرساليات", 1)
doc.add_paragraph("يعرض هذا القسم جميع النماذج التي تم إرسالها مع حالة كل إرسالية.")

add_table(
    ["الحالة", "الوصف"],
    [
        ["🟢 مُرسلة", "تم إرسالها بنجاح والسيرفر استلمها"],
        ["🟡 قيد المراجعة", "بانتظار موافقة المشرف"],
        ["✅ مُعتمدة", "تمت الموافقة عليها من المشرف"],
        ["🔴 مرفوضة", "تم رفضها مع سبب الرفض"],
        ["⏳ في الانتظار", "لم تُزامن بعد (بدون إنترنت)"],
    ]
)

add_heading_styled("الموافقة والرفض (للمشرفين)", 2)
add_step(1, "اضغط على الإرسالية لعرض التفاصيل")
add_step(2, "راجع البيانات والموقع والصور المرفقة")
add_step(3, "اضغط 'اعتماد' للقبول أو 'رفض' مع كتابة السبب")

doc.add_page_break()

# ═══ SECTION 6 ═══
add_heading_styled("6. لوحة التحليلات", 1)
doc.add_paragraph("توفر التحليلات نظرة شاملة على أداء حملات التطعيم مع رسوم بيانية وتصدير التقارير.")

add_table(
    ["المؤشر", "الوصف"],
    [
        ["إجمالي التطعيمات", "عدد التطعيمات حسب الفترة المحددة"],
        ["التوزيع الجغرافي", "التطعيمات حسب المحافظة والمديرة"],
        ["نسبة الإنجاز", "نسبة التطعيمات المكتملة من المستهدف"],
        ["حسب نوع التطعيم", "توزيع التطعيمات حسب النوع"],
        ["حسب العمر", "توزيع الفئات العمرية المستهدفة"],
    ]
)

add_heading_styled("تصدير التقارير", 2)
add_step(1, "اضغط على زر 'تصدير' في أعلى شاشة التحليلات")
add_step(2, "اختر صيغة التصدير: PDF أو CSV")
add_step(3, "حدد الفترة الزمنية والمحافظة (إن أردت)")
add_step(4, "سيتم تحميل الملف تلقائياً")

doc.add_page_break()

# ═══ SECTION 7 ═══
add_heading_styled("7. المساعد الذكي AI", 1)
doc.add_paragraph("مساعد ذكي يدعم اللغة العربية للإجابة على استفساراتك حول بيانات التطعيم وتحليلها.")

add_step(1, "افتح قسم 'الذكاء الاصطناعي' من القائمة")
add_step(2, "اكتب سؤالك باللغة العربية")
add_step(3, "المساعد سيرد بالأرقام والتحليل")

add_info_box("أمثلة: 'كم عدد التطعيمات في تعز هذا الأسبوع؟' — 'ما نسبة الإنجاز في كل المحافظات؟' — 'أي المديريات تحتاج دعماً إضافياً؟'")

doc.add_page_break()

# ═══ SECTION 8 ═══
add_heading_styled("8. الخريطة التفاعلية", 1)
doc.add_paragraph("تعرض الخريطة مواقع الإرساليات الجغرافية مع إمكانية التكبير والتصغير والتصفية.")

add_table(
    ["الميزة", "الوصف"],
    [
        ["📍 علامات المواقع", "نقطة زرقاء لكل إرسالية بناءً على GPS"],
        ["🔢 التجميع", "تجمع النقاط القريبة عند التصغير"],
        ["🌡️ خريطة الحرارة", "ألوان توضح كثافة التطعيمات"],
        ["🔍 التصفية", "تصفية حسب المحافظة ونوع التطعيم"],
        ["📋 تفاصيل النقطة", "اضغط على أي نقطة لعرض التفاصيل"],
    ]
)

doc.add_page_break()

# ═══ SECTION 9 ═══
add_heading_styled("9. إدارة المستخدمين (للمدير)", 1)
doc.add_paragraph("متاح فقط لدور admin و central. يمكنك إدارة حسابات المستخدمين وأدوارهم.")

add_step(1, "اذهب إلى القائمة ← إدارة المستخدمين")
add_step(2, "اضغط على 'إضافة مستخدم' (+)")
add_step(3, "املأ البيانات: الاسم، البريد، كلمة المرور، الدور")
add_step(4, "اختر المحافظة/المديرة حسب الدور")
add_step(5, "اضغط 'حفظ'")

doc.add_page_break()

# ═══ SECTION 10 ═══
add_heading_styled("10. إدارة النماذج (للمدير)", 1)
doc.add_paragraph("يمكن لمدير النظام إنشاء وتعديل النماذج الديناميكية التي يستخدمها مدخلو البيانات.")

add_step(1, "اذهب إلى القائمة ← إدارة النماذج")
add_step(2, "اضغط 'إنشاء نموذج جديد'")
add_step(3, "أدخل اسم النموذج ووصفه")
add_step(4, "أضف الحقول: نص، رقم، تاريخ، قائمة منسدلة، صورة، موقع GPS")
add_step(5, "اضغط 'حفظ ونشر'")

add_info_box("أنواع الحقول المتاحة: نص قصير، نص طويل، رقم، تاريخ، قائمة منسدلة، اختيار متعدد، صورة، موقع GPS، توقيع، ملف مرفق")

doc.add_page_break()

# ═══ SECTION 11 ═══
add_heading_styled("11. العمل بدون إنترنت", 1)
doc.add_paragraph("المنصة مصممة بالكامل للعمل Offline-First — العمل بدون إنترنت هو الحالة الطبيعية.")

add_step(1, "الحفظ المحلي أولاً — كل نموذج يُحفظ فوراً على جهازك")
add_step(2, "طابور المزامنة — النماذج تنتظر بأولويات (التطعيمات أولاً)")
add_step(3, "المزامنة التلقائية — عند عودة الإنترنت تُرسل كل 5 دقائق")
add_step(4, "إعادة المحاولة — 10 ثواني ← 30 ثانية ← 90 ثانية ← 5 دقائق ← 15 دقيقة")
add_step(5, "حل التعارضات — إذا عُدّل نفس السجل على جهازين يُدمج تلقائياً")

add_table(
    ["الرمز", "الحالة", "الوصف"],
    [
        ["🟢", "متصل", "كل شيء متزامن"],
        ["🟡", "متصل مع تأخير", "سجلات في الانتظار"],
        ["🔴", "غير متصل", "البيانات محفوظة محلياً"],
    ]
)

add_success_box("ضمان عدم فقدان البيانات: حتى لو انطفأ الجهاز — بياناتك آمنة.")

doc.add_page_break()

# ═══ SECTION 12 ═══
add_heading_styled("12. نظام الصلاحيات (RBAC)", 1)
doc.add_paragraph("المنصة تستخدم نظام صلاحيات هرمي بـ 5 مستويات.")

add_table(
    ["الدور", "المستوى", "الصلاحيات"],
    [
        ["مدير النظام (admin)", "5", "كل شيء: إدارة المستخدمين والنماذج"],
        ["مشرف مركزي (central)", "4", "رؤية كل البيانات وإدارة النماذج"],
        ["مشرف محافظة (governorate)", "3", "رؤية بيانات محافظته والموافقة"],
        ["مشرف مديرية (district)", "2", "رؤية بيانات مديريته فقط"],
        ["مدخل بيانات (data_entry)", "1", "ملء النماذج ورؤية بياناته فقط"],
    ]
)

add_warn_box("مدخل البيانات لا يستطيع رؤية إرساليات الآخرين أو الموافقة/الرفض.")

doc.add_page_break()

# ═══ SECTION 13 ═══
add_heading_styled("13. الإشعارات", 1)
doc.add_paragraph("تستلم إشعارات عن الأحداث المهمة في المنصة.")

add_table(
    ["النوع", "الوصف"],
    [
        ["📤 تأكيد الإرسال", "تم استلام نموذجك بنجاح"],
        ["✅ اعتماد", "تمت الموافقة على إرساليتك"],
        ["❌ رفض", "تم رفض إرساليتك — اضغط لرؤية السبب"],
        ["⚠️ نقص تجهيزات", "تم الإبلاغ عن نقص في منطقة قريبة"],
        ["📢 إشعار من النظام", "تحديثات وإرشادات من الإدارة"],
    ]
)

doc.add_page_break()

# ═══ SECTION 14 ═══
add_heading_styled("14. استكشاف الأخطاء وحلها", 1)

add_heading_styled("لا أستطيع تسجيل الدخول", 2)
add_step(1, "تأكد من صحة البريد الإلكتروني وكلمة المرور")
add_step(2, "تأكد من اتصالك بالإنترنت")
add_step(3, "تواصل مع مدير النظام إذا استمرت المشكلة")

add_heading_styled("النماذج لا تُرسل", 2)
add_step(1, "تحقق من مؤشر الاتصال — إذا 🔴 فالبيانات محفوظة وستُرسل عند عودة الإنترنت")
add_step(2, "تأكد من ملء جميع الحقول المطلوبة")
add_step(3, "تحقق من تفعيل GPS إذا كان مطلوباً")

add_heading_styled("التطبيق بطيء", 2)
add_step(1, "أغلق التطبيق وأعد فتحه")
add_step(2, "امسح ذاكرة التخزين المؤقت")
add_step(3, "تأكد من وجود مساحة كافية على جهازك")

add_info_box("الدعم الفني: support@epi-supervisor.com")

# ═══ FOOTER ═══
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\nمنصة مشرف EPI — الإصدار 1.0.0\nأبريل 2026\n© جميع الحقوق محفوظة")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_dir = "/root/.openclaw/workspace/EPI-Supervisor/docs/user-guide"
doc.save(os.path.join(output_dir, "EPI_Supervisor_User_Guide.docx"))
print("✅ DOCX generated successfully!")
